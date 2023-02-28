from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

# Open document
document = Document('nmaptemplate.docx')

# Get table
table = document.tables[0]

# Get first cell
cell = table.cell(8, 1)

# Get paragraph in cell
paragraph = cell.paragraphs[0]

# Get formatting from cell
font = paragraph.runs[0].font
color = paragraph.runs[0].font.color.rgb


# Update text and formatting
paragraph.clear()
paragraph.add_run('shyam')
paragraph.runs[0].font.name = font.name
paragraph.runs[0].font.size = font.size
paragraph.runs[0].font.bold = font.bold
paragraph.runs[0].font.italic = font.italic
paragraph.runs[0].font.underline = font.underline
paragraph.runs[0].font.color.rgb = color

# Save document
document.save('example.docx')
