import json
from docx import Document
from docx.shared import RGBColor

def update_table_cells(document_name, dict_file):
    # Open document
    document = Document(document_name)

    # Get table
    table = document.tables[0]

    # Open dictionary file
    with open(dict_file, 'r') as f:
        cell_values = json.load(f)

    # Update cells
    for cell_ref, cell_data in cell_values.items():
        cell_text = cell_data['value']

        # Check if append or overwrite flag is set
        if cell_data['mode'].lower() == 'a':
            append_data = True
        else:
            append_data = False

        # Convert cell_ref to integer values
        row_idx, col_idx = map(int, cell_ref.split(','))

        # Get cell
        cell = table.cell(row_idx, col_idx)

        # Get paragraph in cell
        paragraph = cell.paragraphs[0]

        # Get formatting from cell
        font = paragraph.runs[0].font
        color = paragraph.runs[0].font.color.rgb

        # Update text and formatting
        if append_data:
            paragraph.add_run(f"\n{cell_text}")
        else:
            paragraph.clear()
            paragraph.add_run(cell_text)
        paragraph.runs[0].font.name = font.name
        paragraph.runs[0].font.size = font.size
        paragraph.runs[0].font.bold = font.bold
        paragraph.runs[0].font.italic = font.italic
        paragraph.runs[0].font.underline = font.underline
        paragraph.runs[0].font.color.rgb = color

    # Save document
    document.save(document_name)
